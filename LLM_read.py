import os
import re
from docx import Document
import pandas as pd
import yaml


class LLMAnalyzer:
    # 定义dic
    def __init__(self):
        self.prompts_match_dic = {
            '因子组成': '配置逻辑',
            '因子挖掘': '配置逻辑',
            '模型种类': '配置逻辑',
            '选股范围': '配置逻辑+持仓特征',
            '交易频率': '交易风格',
            '持股数量': '持仓特征'
        }

        self.dic_content_extraction = {'配置逻辑': r'配置逻辑(.*?)持仓特征',
                                       '公司概况': r'公司概况(.*?)股东情况',
                                       '交易风格': r'交易风格(.*?)容量',
                                       '持仓特征': r'持仓特征(.*?)交易风格',
                                       '配置逻辑+持仓特征': r'配置逻辑(.*?)容量'}

    # 获取config的函数
    @staticmethod
    def get_prompt_configs(prompt_class):
        with open('configs/prompts.yaml', 'r', encoding='utf-8') as file:
            config = yaml.safe_load(file)
        return config[prompt_class].replace('\n', '')

    @staticmethod
    def get_text_configs(text_class):
        with open('configs/content_extraction.yaml', 'r', encoding='utf-8') as file:
            config = yaml.safe_load(file)
        return config[text_class].replace('\n', '').replace('"', '')

    # 把prompt中的xxxxxx替换成尽调报告中的对应内容
    @staticmethod
    def get_prompt(text, prompt):
        return prompt.replace('xxxxxx', text)

    # 生成百川所需的特定的prompt
    def get_prompt_baichuan(self, text, prompt):
        messages = [{"role": "user", "content": self.get_prompt(text, prompt)}]
        return messages

    # 从指定目录中获取尽调报告的docx文档，然后提取出“配置逻辑”的内容
    @staticmethod
    def extract_content_from_docx(folder_path, regex_pattern):
        extracted_contents = []
        # 遍历文件夹中的所有文件
        for file in os.listdir(folder_path):
            if file.endswith(".docx"):
                # 构造完整的文件路径
                file_path = os.path.join(folder_path, file)
                # 打开Word文档
                doc = Document(file_path)
                for para in doc.paragraphs:
                    if "交易特征" in para.text:
                        para.text = para.text.replace("交易特征", "交易风格")
                # 读取文档中的所有段落
                full_text = '\n'.join([para.text for para in doc.paragraphs])
                # 使用正则表达式匹配内容
                matches = re.search(regex_pattern, full_text, re.DOTALL)
                matches = matches.group().replace('\n', '')
                if matches:
                    extracted_contents.append(matches)
        return extracted_contents

    @staticmethod
    def get_name_list(folder_path):
        name_list = []
        for file in os.listdir(folder_path):
            if file.endswith(".docx"):
                name_list.append(file.replace(".docx", ""))
        return name_list

    def get_answer_from_docx_glm(self, extracted_contents, prompt1, name_list, model, tokenizer):
        df_dic = {
            "私募名称": name_list,
            "因子组成": [],
        }
        temporary_list = []
        for text in extracted_contents:
            response, his = model.chat(tokenizer, self.get_prompt(text, prompt1), history=[], temperature=0.1)
            temporary_list.append(response)
            df_dic['因子组成'].append(response)
        return pd.DataFrame(df_dic), temporary_list

    def get_answer_template(self, extracted_contents, prompt, model):
        response_list = []
        count = 0
        for text in extracted_contents:
            response = model(self.get_prompt_baichuan(text, prompt))
            response_list.append(response['response'])
            print(count)
            count += 1
        return response_list

    def get_answer_from_docx_baichuan(self, yaml_list, model, doc_path):
        name_list = self.get_name_list(doc_path)
        # 读取 n 个yaml_list中的元素，并创建字典
        df_dic = {
            "私募名称": name_list
        }

        for prompt_class in yaml_list:
            df_dic[prompt_class] = []
            prompt = self.get_prompt_configs(prompt_class)
            prompt_key = self.prompts_match_dic[prompt_class]
            # 按照 yaml_list中的要求，提取prompt并使用大语言模型
            extracted_contents = self.extract_content_from_docx("尽调报告", self.dic_content_extraction[prompt_key])
            responses = self.get_answer_template(extracted_contents, prompt, model)
            df_dic[prompt_class] = responses

        return pd.DataFrame(df_dic)
